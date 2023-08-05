import PySimpleGUI as sg
from docx import Document
import os
import sys
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
import pyautogui
import subprocess

def save_data(data):
    document = Document()

    # Set font and size
    font_name = "Times New Roman"
    font_size = 12

    # Add the truck image
    truck_image_path = 'C:/Users/14842/Documents/App/trucking.png'
    document.add_picture(truck_image_path)

    # Create a custom style for the body text
    body_text_style = document.styles.add_style("BodyText", WD_STYLE_TYPE.PARAGRAPH)
    body_text_style.font.name = font_name
    body_text_style.font.size = Pt(font_size)

    # Add Move # and Date headers
    move_number = data["-MOVE-"]
    move_date = data["-DATE-"]
    header_paragraph = document.add_paragraph()

    # Add Move # and its input
    move_text = f"Move # "
    move_run = header_paragraph.add_run(move_text)
    move_run.font.name = font_name
    move_run.font.size = Pt(font_size)

    # Add bold input value for Move
    move_input_text = str(move_number)
    move_input_run = header_paragraph.add_run(move_input_text)
    move_input_run.font.name = font_name
    move_input_run.font.size = Pt(font_size)
    move_input_run.bold = True

    # Add spacing between Move # and Date:
    header_paragraph.add_run("                                                                  " * 1)

    # Add Date: and its input
    date_text = f"Date: "
    date_run = header_paragraph.add_run(date_text)
    date_run.font.name = font_name
    date_run.font.size = Pt(font_size)

    # Add bold input value for Date
    date_input_text = str(move_date)
    date_input_run = header_paragraph.add_run(date_input_text)
    date_input_run.font.name = font_name
    date_input_run.font.size = Pt(font_size)
    date_input_run.bold = True


    document.add_paragraph("_________________________________________________________________________________________________________")

    # Create a table with two columns
    table = document.add_table(rows=11, cols=2)
    table.style = "Table Grid"

    # Set column widths
    table.columns[0].width = Cm(8)
    table.columns[1].width = Cm(8)

    # Add Description and Amount Due labels as headers
    headers = table.rows[0].cells
    headers[0].text = "Description:"
    headers[1].text = "Amount Due:"
    for cell in headers:
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = font_name
        cell.paragraphs[0].runs[0].font.size = Pt(font_size)

    # Add Line Haul 1/2 inputs and their values
    line_haul_1 = data.get("-LINE_HAUL-1")
    line_haul_2 = data.get("-LINE_HAUL-2")
    table.cell(1, 0).text = f"Line Haul: {line_haul_1}"
    table.cell(1, 1).text = f"${line_haul_2}" if line_haul_2 else "$0"

    # Add Lumper Fee 1/2 inputs and their values
    lumper_fee_1 = data.get("-LUMPER_FEE-1")
    lumper_fee_2 = data.get("-LUMPER_FEE-2")                       
    table.cell(2, 0).text = f"Lumper Fee: {lumper_fee_1}"
    table.cell(2, 1).text = f"${lumper_fee_2}" if lumper_fee_2 else "$0"

    # Add Stop Off 1/2 inputs and their values
    stop_off_1 = data.get("-STOP_OFF-1")
    stop_off_2 = data.get("-STOP_OFF-2")
    table.cell(3, 0).text = f"Stop Off: {stop_off_1}"
    table.cell(3, 1).text = f"${stop_off_2}" if stop_off_2 else "$0"

    # Add Returns 1/2 inputs and their values
    returns_1 = data.get("-RETURNS-1")
    returns_2 = data.get("-RETURNS-2")
    table.cell(4, 0).text = f"Returns: {returns_1}"
    table.cell(4, 1).text = f"${returns_2}" if returns_2 else "$0"

    # Add Detention 1/2 inputs and their values
    detention_1 = data.get("-DETENTION-1")
    detention_2 = data.get("-DETENTION-2")
    table.cell(5, 0).text = f"Detention: {detention_1}"
    table.cell(5, 1).text = f"${detention_2}" if detention_2 else "$0"

    # Add Additional labels and inputs
    table.cell(6, 0).text = "Additional 1:"
    table.cell(6, 0).text = data.get('-ADDITIONAL-1', '0')
    table.cell(6, 1).text = "Additional 2:"
    table.cell(6, 1).text = "$" + str(data.get('-ADDITIONAL-2', '0') if data.get('-ADDITIONAL-2') else '0')
    table.cell(7, 0).text = "Additional 3:"
    table.cell(7, 0).text = data.get('-ADDITIONAL-3', '0')
    table.cell(7, 1).text = "Additional 4:"
    table.cell(7, 1).text = "$" + str(data.get('-ADDITIONAL-4', '0') if data.get('-ADDITIONAL-4') else '0')
    table.cell(8, 0).text = "Additional 5:"
    table.cell(8, 0).text = data.get('-ADDITIONAL-5', '0')
    table.cell(8, 1).text = "Additional 6:"
    table.cell(8, 1).text = "$" + str(data.get('-ADDITIONAL-6', '0') if data.get('-ADDITIONAL-6') else '0')
    table.cell(9, 0).text = "Additional 7:"
    table.cell(9, 0).text = data.get('-ADDITIONAL-7', '0')
    table.cell(9, 1).text = "Additional 8:"
    table.cell(9, 1).text = "$" + str(data.get('-ADDITIONAL-8', '0') if data.get('-ADDITIONAL-8') else '0')
    table.cell(10, 0).text = "Additional 9:"
    table.cell(10, 0).text = data.get('-ADDITIONAL-9', '0')
    table.cell(10, 1).text = "Additional 10:"
    table.cell(10, 1).text = "$" + str(data.get('-ADDITIONAL-10', '0') if data.get('-ADDITIONAL-10') else '0')


    document.add_paragraph("_________________________________________________________________________________________________________")

    # Add Lumper Fee Paid input and its value
    lumper_fee_paid = data.get("-LUMPER_FEE_PAID-")
    if lumper_fee_paid is not None and lumper_fee_paid != "":
        lumper_fee_paid = int(lumper_fee_paid)
    else:
        lumper_fee_paid = 0

    if lumper_fee_paid != -20:
        lumper_fee_paid = 0

    lumper_paragraph = document.add_paragraph(style=body_text_style)
    lumper_run = lumper_paragraph.add_run(f"Com Data Fee: ${lumper_fee_paid}")
    lumper_run.font.name = font_name
    lumper_run.font.size = Pt(font_size)
    lumper_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


   # Add Total Invoice Amount and its value
    invoice_amount = calculate_invoice_amount(data)
    amount_paragraph = document.add_paragraph(style=body_text_style)
    amount_run = amount_paragraph.add_run("Total Invoice Amount:")
    amount_run.bold = True
    amount_run.font.name = font_name
    amount_run.font.size = Pt(font_size)

    amount_value_run = amount_paragraph.add_run(f" ${invoice_amount}")
    amount_value_run.bold = True
    amount_value_run.font.name = font_name
    amount_value_run.font.size = Pt(font_size)

    amount_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Get the directory of the executable file
    if getattr(sys, 'frozen', False):
        script_dir = os.path.dirname(sys.executable)
    else:
        script_dir = os.path.dirname(os.path.abspath(__file__))

    # Construct the file path
    file_path = os.path.join(script_dir, 'Invoice_Pay.docx')
    document.save(file_path)

    # Open the document
    os.startfile(file_path)

    # Wait for the application to open
    sg.popup_quick_message("Opening the document... Please wait.")

    # Simulate key presses to print the document
    pyautogui.hotkey('ctrl', 'p')
    pyautogui.press('enter')


def calculate_invoice_amount(values):
    line_haul_2 = float(values["-LINE_HAUL-2"]) if values["-LINE_HAUL-2"] else 0
    lumper_fee_2 = float(values["-LUMPER_FEE-2"]) if values["-LUMPER_FEE-2"] else 0
    stop_off_2 = float(values["-STOP_OFF-2"]) if values["-STOP_OFF-2"] else 0
    returns_2 = float(values["-RETURNS-2"]) if values["-RETURNS-2"] else 0
    detention_2 = float(values["-DETENTION-2"]) if values["-DETENTION-2"] else 0
    additional_values = [
        float(values["-ADDITIONAL-2"]) if values["-ADDITIONAL-2"] else 0,
        float(values["-ADDITIONAL-4"]) if values["-ADDITIONAL-4"] else 0,
        float(values["-ADDITIONAL-6"]) if values["-ADDITIONAL-6"] else 0,
        float(values["-ADDITIONAL-8"]) if values["-ADDITIONAL-8"] else 0,
        float(values["-ADDITIONAL-10"]) if values["-ADDITIONAL-10"] else 0
    ]
    invoice_amount = line_haul_2 + lumper_fee_2 + stop_off_2 + returns_2 + detention_2 + sum(additional_values)

    lumper_fee_paid = values["-LUMPER_FEE_PAID-"]
    if lumper_fee_paid == "-20":
        invoice_amount -= 20    

    invoice_amount_formatted = "{:.2f}".format(invoice_amount)
    return invoice_amount_formatted

def main():
    icon_path = r"C:/Users/14842/Documents/App/dist/truck3.ico"
    sg.set_options(icon=icon_path)  # Set the program icon

    truck_image_path =  'C:/Users/14842/Documents/App/truck.png'
    truck_image = sg.Image(truck_image_path, size=(400, 300))
    
    header_layout = [
        [
            sg.Frame('', [
                [sg.Text('VASQUEZ', font=('Helvetica', 30, 'bold'), text_color='#333333')],
                [sg.Text('TRUCKING LLC', font=('Helvetica', 30, 'bold'), text_color='#333333')],
                [sg.Text('48 Winding Brook Dr. Sinking Spring, PA 19608', font=('Helvetica', 13), text_color='#333333')],
                [sg.Text('484-269-6029', font=('Helvetica', 13), text_color='#333333')]
            ], background_color='#444444', pad=(0, 0))
        ]
    ]

    truck_image = sg.Image('C:/Users/14842/Documents/App/truck.png', size=(400, 200))

    # Define the layout
    layout = [
        [
            sg.Column(header_layout, element_justification='left'),
            sg.Column([[truck_image]], element_justification='right'), 
        ],
        [
            sg.Text("Move #"),
            sg.Input(key="-MOVE-", size=(10, 1)),
            sg.Text("Date:"),
            sg.Input(key="-DATE-", enable_events=True, readonly=True, size=(15, 1)),
            sg.CalendarButton("Select Date", target="-DATE-", key="-DATE_BTN-", format="%Y-%m-%d", size=(12, 1))
        ],
        [
            sg.Text("Description:", size=(20, 1), justification='right', font=('Helvetica', 11, 'bold'), text_color='black'),
            sg.Text("Amount Due:", size=(35, 1), justification='right', font=('Helvetica', 11, 'bold'), text_color='black')
        ],
        [
            sg.Text("Line Haul:", size=(10, 1)),
            sg.Input(key="-LINE_HAUL-1", justification="left"),
            sg.Input(key="-LINE_HAUL-2", justification="left")
        ],
        [
            sg.Text("Lumper Fee:", size=(10, 1)),
            sg.Input(key="-LUMPER_FEE-1", justification="left"),
            sg.Input(key="-LUMPER_FEE-2", justification="left")
        ],
        [
            sg.Text("Stop Off:", size=(10, 1)),
            sg.Input(key="-STOP_OFF-1", justification="left"),
            sg.Input(key="-STOP_OFF-2", justification="left")
        ],
        [
            sg.Text("Returns:", size=(10, 1)),
            sg.Input(key="-RETURNS-1", justification="left"),
            sg.Input(key="-RETURNS-2", justification="left")
        ],
        [
            sg.Text("Detention:", size=(10, 1)),
            sg.Input(key="-DETENTION-1", justification="left"),
            sg.Input(key="-DETENTION-2", justification="left")
        ],
        [
    sg.Text("Additional:", size=(10, 1)),
    sg.Input(key="-ADDITIONAL-1", justification="left"),
    sg.Input(key="-ADDITIONAL-2", justification="left")
],
[
    sg.Text("", size=(10, 1)),
    sg.Input(key="-ADDITIONAL-3", justification="left"),
    sg.Input(key="-ADDITIONAL-4", justification="left")
],
[
    sg.Text("", size=(10, 1)),
    sg.Input(key="-ADDITIONAL-5", justification="left"),
    sg.Input(key="-ADDITIONAL-6", justification="left")
],
[
    sg.Text("", size=(10, 1)),
    sg.Input(key="-ADDITIONAL-7", justification="left"),
    sg.Input(key="-ADDITIONAL-8", justification="left")
],
[
    sg.Text("", size=(10, 1)),
    sg.Input(key="-ADDITIONAL-9", justification="left"),
    sg.Input(key="-ADDITIONAL-10", justification="left")
],


        [
            sg.Text("Com Data Fee:", size=(15, 1)),
            sg.Combo(["", "-20"], key="-LUMPER_FEE_PAID-", enable_events=True, readonly=True, size=(20, 1))
        ],
        [
            sg.Button("Calculate", key="-CALCULATE-", button_color=("white", "green")),
        sg.Button("Clear", key="-CLEAR-", button_color=("white", "red")),
        sg.Button("Exit", key="-EXIT-"),
        ]
    ]

    # Create the window
    window = sg.Window("Invoice Pay", layout)

    # Event loop
    while True:
        event, values = window.read()

        # Handle events
        if event == "-CALCULATE-":
            invoice_amount = calculate_invoice_amount(values)
            sg.popup(f"Invoice Amount: ${invoice_amount}")
            save_data(values)  # Save the data to a Word document
            window["-MOVE-"].SetFocus()  # Return focus to "Move #" input
    
        # Simulate key presses to print the document
            pyautogui.hotkey('ctrl', 'p')
            pyautogui.press('enter')

        elif event == "-CLEAR-":
            window["-MOVE-"].SetFocus()  # Return focus to "Move #" input
            window["-MOVE-"].update("")  # Clear the "Move #" input
            window["-DATE-"].update("")  # Clear the date input field
            window["-DATE_BTN-"].update("Select Date")  # Restore the button text
            for key in values:
                if key not in ["-DATE-", "-DATE_BTN-"]:  # Skip updating the button text and date field
                    window[key].update("")  # Clear all other input fields

        elif event == "-EXIT-" or event == sg.WINDOW_CLOSED:
            break

        elif event == "-DATE_BTN-":
            window["-DATE-"].update(values["-DATE_BTN-"].strftime("%Y-%m-%d"))

    window.close()

if __name__ == "__main__":
    main()